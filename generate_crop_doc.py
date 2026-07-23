"""
Generate a Word document (.docx) documenting which crops can be recommended
by the AI Farming Project's Crop Advisor module.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Crop_Recommendations.docx"
)


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "2E7D32")
        borders.append(b)
    tblPr.append(borders)


def header_row(table, row_idx, color="1B5E20"):
    for cell in table.rows[row_idx].cells:
        set_cell_shading(cell, color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)


def create_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(5):
        doc.add_paragraph("")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AI Agriculture Project")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Plant Disease Detection & Smart Farming")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph("")
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = line.add_run("_" * 50)
    r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph("")
    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dt.add_run("CROP RECOMMENDATION GUIDE")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph("")
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = desc.add_run(
        "Complete documentation of all crops supported by the\n"
        "AGRI-X AI Crop Advisor module, including growing conditions,\n"
        "yield estimates, market prices, and suitability scoring criteria."
    )
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS
    # ================================================================
    h = doc.add_heading("Table of Contents", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    toc = [
        "1. Introduction",
        "2. Crop Recommendation Overview",
        "3. Supported Crops - Summary Table",
        "4. Detailed Crop Profiles",
        "   4.1 Rice",
        "   4.2 Wheat",
        "   4.3 Tomato",
        "   4.4 Potato",
        "   4.5 Bell Pepper (Capsicum)",
        "   4.6 Cotton",
        "   4.7 Sugarcane",
        "   4.8 Maize (Corn)",
        "   4.9 Soybean",
        "   4.10 Mustard",
        "   4.11 Onion",
        "   4.12 Chilli",
        "5. Suitability Scoring Algorithm",
        "6. Season Classification",
        "7. Soil Types Supported",
        "8. Revenue & Profit Estimation",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(3)
        p.runs[0].font.size = Pt(10.5)
        p.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_page_break()

    # ================================================================
    # 1. INTRODUCTION
    # ================================================================
    h = doc.add_heading("1. Introduction", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "The AGRI-X AI Crop Advisor is a smart recommendation module integrated into the "
        "AI Agriculture Project. It suggests the most suitable crops for a farmer based on "
        "real-time environmental conditions, soil type, geographic location, season, and "
        "available land area."
    )
    doc.add_paragraph(
        "The system currently supports 12 major crops commonly grown across India. Each crop "
        "is evaluated against the farmer's conditions using a multi-factor suitability scoring "
        "algorithm that considers temperature, humidity, rainfall, soil type, and growing season."
    )
    doc.add_paragraph(
        "In addition to crop selection, the module provides yield estimates, market price ranges, "
        "Minimum Support Prices (MSP) where applicable, input cost estimates, and projected "
        "revenue and profit calculations."
    )

    # ================================================================
    # 2. CROP RECOMMENDATION OVERVIEW
    # ================================================================
    h = doc.add_heading("2. Crop Recommendation Overview", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "The recommendation engine works in the following steps:"
    )
    steps = [
        ("Step 1: Input Collection", "Farmer provides land area, soil type, and optionally adjusts the season. "
         "Temperature, humidity, and estimated annual rainfall are auto-filled from live weather data based on the farmer's GPS location."),
        ("Step 2: Crop Scoring", "Each of the 12 crops in the database is scored (0-100) based on how well the "
         "current conditions match that crop's ideal growing requirements."),
        ("Step 3: Ranking", "Crops are ranked by suitability score (highest first). Crops scoring below 20 are excluded."),
        ("Step 4: Yield & Revenue", "For recommended crops, the system estimates expected yield (tonnes/acre), "
         "revenue, input costs, and net profit based on the farmer's land area."),
        ("Step 5: Presentation", "Results are displayed as interactive cards with detailed suitability breakdowns, "
         "plus a comparative revenue analysis chart."),
    ]
    for title, detail in steps:
        p = doc.add_paragraph()
        r = p.add_run(title + ": ")
        r.bold = True
        r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        p.add_run(detail)

    # ================================================================
    # 3. SUPPORTED CROPS SUMMARY
    # ================================================================
    doc.add_page_break()
    h = doc.add_heading("3. Supported Crops - Summary Table", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "The following table lists all 12 crops in the recommendation database with their key parameters:"
    )

    # Main summary table
    headers = ["#", "Crop", "Season", "Temp (C)", "Rainfall (mm)", "Water Need", "Duration (days)"]
    crops_summary = [
        ("1",  "Rice",        "Kharif",       "20-35",  "1000-2000", "High",      "120"),
        ("2",  "Wheat",       "Rabi",         "10-25",  "400-700",   "Moderate",  "135"),
        ("3",  "Tomato",      "Kharif/Rabi",  "18-30",  "400-600",   "Moderate",  "100"),
        ("4",  "Potato",      "Rabi",         "15-25",  "500-800",   "Moderate",  "100"),
        ("5",  "Bell Pepper", "Kharif/Rabi",  "18-28",  "600-1200",  "Moderate",  "90"),
        ("6",  "Cotton",      "Kharif",       "25-35",  "500-1000",  "Moderate",  "165"),
        ("7",  "Sugarcane",   "Kharif/Rabi",  "20-35",  "1500-2500", "Very High", "330"),
        ("8",  "Maize",       "Kharif",       "20-30",  "500-1000",  "Moderate",  "100"),
        ("9",  "Soybean",     "Kharif",       "20-30",  "600-1000",  "Moderate",  "100"),
        ("10", "Mustard",     "Rabi",         "10-25",  "300-500",   "Low",       "120"),
        ("11", "Onion",       "Rabi",         "15-25",  "350-600",   "Moderate",  "130"),
        ("12", "Chilli",      "Kharif",       "20-35",  "500-1200",  "Moderate",  "150"),
    ]

    table = doc.add_table(rows=len(crops_summary) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    for j, h_text in enumerate(headers):
        table.rows[0].cells[j].text = h_text
    header_row(table, 0)

    for i, row_data in enumerate(crops_summary):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val
        # Alternate row shading
        if i % 2 == 1:
            for cell in table.rows[i + 1].cells:
                set_cell_shading(cell, "F1F8E9")

    doc.add_paragraph("")

    # Yield & Economics summary table
    h2 = doc.add_heading("3.1 Yield & Economics Summary", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    econ_headers = ["Crop", "Yield (t/acre)", "Price (Rs/quintal)", "MSP (Rs/q)", "Cost (Rs/acre)"]
    econ_data = [
        ("Rice",        "1.0 - 1.5",  "2,000 - 2,400",  "2,183",  "18,000"),
        ("Wheat",       "1.2 - 1.8",  "2,100 - 2,500",  "2,275",  "16,000"),
        ("Tomato",      "8.0 - 15.0", "800 - 2,500",    "N/A",    "45,000"),
        ("Potato",      "6.0 - 10.0", "600 - 1,500",    "N/A",    "35,000"),
        ("Bell Pepper", "4.0 - 8.0",  "1,500 - 4,000",  "N/A",    "55,000"),
        ("Cotton",      "0.6 - 0.8",  "6,000 - 7,500",  "6,620",  "22,000"),
        ("Sugarcane",   "25 - 35",    "290 - 350",       "315",    "40,000"),
        ("Maize",       "1.0 - 1.5",  "1,900 - 2,300",  "2,090",  "14,000"),
        ("Soybean",     "0.4 - 0.6",  "4,200 - 5,000",  "4,600",  "15,000"),
        ("Mustard",     "0.5 - 0.8",  "5,000 - 6,500",  "5,650",  "12,000"),
        ("Onion",       "6.0 - 12.0", "800 - 3,000",    "N/A",    "40,000"),
        ("Chilli",      "1.5 - 3.0",  "8,000 - 15,000", "N/A",    "50,000"),
    ]

    table = doc.add_table(rows=len(econ_data) + 1, cols=len(econ_headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    for j, h_text in enumerate(econ_headers):
        table.rows[0].cells[j].text = h_text
    header_row(table, 0)

    for i, row_data in enumerate(econ_data):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val
        if i % 2 == 1:
            for cell in table.rows[i + 1].cells:
                set_cell_shading(cell, "F1F8E9")

    # ================================================================
    # 4. DETAILED CROP PROFILES
    # ================================================================
    doc.add_page_break()
    h = doc.add_heading("4. Detailed Crop Profiles", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "This section provides a detailed profile for each of the 12 supported crops, "
        "including growing conditions, soil requirements, economic data, and relevance "
        "to the AI disease detection system."
    )

    CROP_PROFILES = [
        {
            "number": "4.1", "name": "Rice", "icon": "Kharif Cereal",
            "description": (
                "Rice is the primary staple food crop of India, cultivated extensively during the "
                "Kharif (monsoon) season. It requires high humidity and abundant water, making it "
                "suitable for regions with heavy rainfall or access to irrigation."
            ),
            "conditions": [
                ("Temperature Range", "20 - 35 C"),
                ("Humidity Range", "60 - 80%"),
                ("Annual Rainfall", "1,000 - 2,000 mm"),
                ("Suitable Soil Types", "Alluvial, Clay, Loamy"),
                ("Growing Season", "Kharif (June - October)"),
                ("Crop Duration", "120 days"),
                ("Water Requirement", "High"),
            ],
            "economics": [
                ("Yield per Acre", "1.0 - 1.5 tonnes"),
                ("Market Price", "Rs 2,000 - 2,400 per quintal"),
                ("MSP (Minimum Support Price)", "Rs 2,183 per quintal"),
                ("Input Cost per Acre", "Rs 18,000"),
            ],
            "disease_link": "Not directly covered by the disease detection model (model covers Pepper, Potato, Tomato).",
        },
        {
            "number": "4.2", "name": "Wheat", "icon": "Rabi Cereal",
            "description": (
                "Wheat is India's second most important cereal crop, grown during the Rabi "
                "(winter) season. It requires cool temperatures during growth and warm, dry "
                "weather during harvesting."
            ),
            "conditions": [
                ("Temperature Range", "10 - 25 C"),
                ("Humidity Range", "50 - 70%"),
                ("Annual Rainfall", "400 - 700 mm"),
                ("Suitable Soil Types", "Alluvial, Loamy, Clay"),
                ("Growing Season", "Rabi (November - March)"),
                ("Crop Duration", "135 days"),
                ("Water Requirement", "Moderate"),
            ],
            "economics": [
                ("Yield per Acre", "1.2 - 1.8 tonnes"),
                ("Market Price", "Rs 2,100 - 2,500 per quintal"),
                ("MSP (Minimum Support Price)", "Rs 2,275 per quintal"),
                ("Input Cost per Acre", "Rs 16,000"),
            ],
            "disease_link": "Not directly covered by the disease detection model.",
        },
        {
            "number": "4.3", "name": "Tomato", "icon": "Kharif/Rabi Vegetable",
            "description": (
                "Tomato is one of the most widely grown vegetable crops in India. It can be "
                "cultivated in both Kharif and Rabi seasons. The AI disease detection model "
                "covers 10 tomato conditions (9 diseases + healthy), making this crop the "
                "most comprehensively monitored crop in the system."
            ),
            "conditions": [
                ("Temperature Range", "18 - 30 C"),
                ("Humidity Range", "50 - 80%"),
                ("Annual Rainfall", "400 - 600 mm"),
                ("Suitable Soil Types", "Loamy, Sandy, Red"),
                ("Growing Season", "Kharif & Rabi (year-round possible)"),
                ("Crop Duration", "100 days"),
                ("Water Requirement", "Moderate"),
            ],
            "economics": [
                ("Yield per Acre", "8.0 - 15.0 tonnes"),
                ("Market Price", "Rs 800 - 2,500 per quintal"),
                ("MSP", "Not applicable (market-driven pricing)"),
                ("Input Cost per Acre", "Rs 45,000"),
            ],
            "disease_link": (
                "Fully covered by the disease detection model with 10 classes: "
                "Bacterial Spot, Early Blight, Late Blight, Leaf Mold, "
                "Septoria Leaf Spot, Spider Mites, Target Spot, Yellow Leaf Curl Virus, "
                "Mosaic Virus, and Healthy."
            ),
        },
        {
            "number": "4.4", "name": "Potato", "icon": "Rabi Tuber",
            "description": (
                "Potato is a major food crop grown during the Rabi season. It is the third "
                "most important food crop globally. The disease detection model covers 3 "
                "potato conditions: Early Blight, Late Blight, and Healthy."
            ),
            "conditions": [
                ("Temperature Range", "15 - 25 C"),
                ("Humidity Range", "60 - 80%"),
                ("Annual Rainfall", "500 - 800 mm"),
                ("Suitable Soil Types", "Loamy, Sandy, Alluvial"),
                ("Growing Season", "Rabi (October - March)"),
                ("Crop Duration", "100 days"),
                ("Water Requirement", "Moderate"),
            ],
            "economics": [
                ("Yield per Acre", "6.0 - 10.0 tonnes"),
                ("Market Price", "Rs 600 - 1,500 per quintal"),
                ("MSP", "Not applicable"),
                ("Input Cost per Acre", "Rs 35,000"),
            ],
            "disease_link": (
                "Covered by the disease detection model with 3 classes: "
                "Early Blight, Late Blight, and Healthy."
            ),
        },
        {
            "number": "4.5", "name": "Bell Pepper (Capsicum)", "icon": "Kharif/Rabi Vegetable",
            "description": (
                "Bell Pepper (Capsicum) is a high-value vegetable crop that can be grown in both "
                "seasons. It requires moderate temperatures and good drainage. The disease detection "
                "model covers Bacterial Spot and Healthy conditions for Bell Pepper."
            ),
            "conditions": [
                ("Temperature Range", "18 - 28 C"),
                ("Humidity Range", "60 - 70%"),
                ("Annual Rainfall", "600 - 1,200 mm"),
                ("Suitable Soil Types", "Loamy, Sandy"),
                ("Growing Season", "Kharif & Rabi"),
                ("Crop Duration", "90 days"),
                ("Water Requirement", "Moderate"),
            ],
            "economics": [
                ("Yield per Acre", "4.0 - 8.0 tonnes"),
                ("Market Price", "Rs 1,500 - 4,000 per quintal"),
                ("MSP", "Not applicable"),
                ("Input Cost per Acre", "Rs 55,000"),
            ],
            "disease_link": (
                "Covered by the disease detection model with 2 classes: "
                "Bacterial Spot and Healthy."
            ),
        },
        {
            "number": "4.6", "name": "Cotton", "icon": "Kharif Cash Crop",
            "description": (
                "Cotton is a major cash crop and the backbone of India's textile industry. "
                "It is grown during the Kharif season and requires warm temperatures with "
                "moderate rainfall."
            ),
            "conditions": [
                ("Temperature Range", "25 - 35 C"),
                ("Humidity Range", "50 - 65%"),
                ("Annual Rainfall", "500 - 1,000 mm"),
                ("Suitable Soil Types", "Black (Regur), Alluvial"),
                ("Growing Season", "Kharif (April - October)"),
                ("Crop Duration", "165 days"),
                ("Water Requirement", "Moderate"),
            ],
            "economics": [
                ("Yield per Acre", "0.6 - 0.8 tonnes"),
                ("Market Price", "Rs 6,000 - 7,500 per quintal"),
                ("MSP (Minimum Support Price)", "Rs 6,620 per quintal"),
                ("Input Cost per Acre", "Rs 22,000"),
            ],
            "disease_link": "Not covered by the disease detection model.",
        },
        {
            "number": "4.7", "name": "Sugarcane", "icon": "Kharif/Rabi Cash Crop",
            "description": (
                "Sugarcane is a long-duration, high-water-demand cash crop used for sugar and "
                "ethanol production. It is one of the most important commercial crops in India "
                "with guaranteed MSP."
            ),
            "conditions": [
                ("Temperature Range", "20 - 35 C"),
                ("Humidity Range", "70 - 85%"),
                ("Annual Rainfall", "1,500 - 2,500 mm"),
                ("Suitable Soil Types", "Alluvial, Loamy, Clay"),
                ("Growing Season", "Kharif & Rabi (perennial, ~11 months)"),
                ("Crop Duration", "330 days"),
                ("Water Requirement", "Very High"),
            ],
            "economics": [
                ("Yield per Acre", "25 - 35 tonnes"),
                ("Market Price", "Rs 290 - 350 per quintal"),
                ("MSP (Fair & Remunerative Price)", "Rs 315 per quintal"),
                ("Input Cost per Acre", "Rs 40,000"),
            ],
            "disease_link": "Not covered by the disease detection model.",
        },
        {
            "number": "4.8", "name": "Maize (Corn)", "icon": "Kharif Cereal",
            "description": (
                "Maize is a versatile cereal used for food, animal feed, and industrial purposes. "
                "It is a fast-growing crop suitable for diverse agro-climatic zones and has "
                "government MSP support."
            ),
            "conditions": [
                ("Temperature Range", "20 - 30 C"),
                ("Humidity Range", "55 - 75%"),
                ("Annual Rainfall", "500 - 1,000 mm"),
                ("Suitable Soil Types", "Loamy, Alluvial, Red"),
                ("Growing Season", "Kharif (June - October)"),
                ("Crop Duration", "100 days"),
                ("Water Requirement", "Moderate"),
            ],
            "economics": [
                ("Yield per Acre", "1.0 - 1.5 tonnes"),
                ("Market Price", "Rs 1,900 - 2,300 per quintal"),
                ("MSP (Minimum Support Price)", "Rs 2,090 per quintal"),
                ("Input Cost per Acre", "Rs 14,000"),
            ],
            "disease_link": "Not covered by the disease detection model.",
        },
        {
            "number": "4.9", "name": "Soybean", "icon": "Kharif Oilseed",
            "description": (
                "Soybean is a major oilseed crop rich in protein, grown extensively in Madhya Pradesh "
                "and Maharashtra. It is a nitrogen-fixing legume beneficial for soil health and "
                "has strong MSP support."
            ),
            "conditions": [
                ("Temperature Range", "20 - 30 C"),
                ("Humidity Range", "60 - 70%"),
                ("Annual Rainfall", "600 - 1,000 mm"),
                ("Suitable Soil Types", "Black (Regur), Loamy"),
                ("Growing Season", "Kharif (June - October)"),
                ("Crop Duration", "100 days"),
                ("Water Requirement", "Moderate"),
            ],
            "economics": [
                ("Yield per Acre", "0.4 - 0.6 tonnes"),
                ("Market Price", "Rs 4,200 - 5,000 per quintal"),
                ("MSP (Minimum Support Price)", "Rs 4,600 per quintal"),
                ("Input Cost per Acre", "Rs 15,000"),
            ],
            "disease_link": "Not covered by the disease detection model.",
        },
        {
            "number": "4.10", "name": "Mustard", "icon": "Rabi Oilseed",
            "description": (
                "Mustard is a key Rabi oilseed crop, especially in Rajasthan and Madhya Pradesh. "
                "It is drought-tolerant and has low water requirements, making it ideal for "
                "semi-arid regions."
            ),
            "conditions": [
                ("Temperature Range", "10 - 25 C"),
                ("Humidity Range", "40 - 60%"),
                ("Annual Rainfall", "300 - 500 mm"),
                ("Suitable Soil Types", "Loamy, Sandy, Alluvial"),
                ("Growing Season", "Rabi (October - February)"),
                ("Crop Duration", "120 days"),
                ("Water Requirement", "Low"),
            ],
            "economics": [
                ("Yield per Acre", "0.5 - 0.8 tonnes"),
                ("Market Price", "Rs 5,000 - 6,500 per quintal"),
                ("MSP (Minimum Support Price)", "Rs 5,650 per quintal"),
                ("Input Cost per Acre", "Rs 12,000"),
            ],
            "disease_link": "Not covered by the disease detection model.",
        },
        {
            "number": "4.11", "name": "Onion", "icon": "Rabi Vegetable",
            "description": (
                "Onion is one of the most commercially important vegetables in India with "
                "highly volatile market prices. It is primarily a Rabi crop but can also be "
                "grown in Kharif in some regions."
            ),
            "conditions": [
                ("Temperature Range", "15 - 25 C"),
                ("Humidity Range", "50 - 70%"),
                ("Annual Rainfall", "350 - 600 mm"),
                ("Suitable Soil Types", "Loamy, Sandy, Alluvial"),
                ("Growing Season", "Rabi (November - March)"),
                ("Crop Duration", "130 days"),
                ("Water Requirement", "Moderate"),
            ],
            "economics": [
                ("Yield per Acre", "6.0 - 12.0 tonnes"),
                ("Market Price", "Rs 800 - 3,000 per quintal"),
                ("MSP", "Not applicable (market-driven)"),
                ("Input Cost per Acre", "Rs 40,000"),
            ],
            "disease_link": "Not covered by the disease detection model.",
        },
        {
            "number": "4.12", "name": "Chilli", "icon": "Kharif Spice",
            "description": (
                "Chilli is a high-value spice crop with strong domestic and export demand. "
                "India is the world's largest producer and exporter of chilli. It has high "
                "input costs but also high potential returns."
            ),
            "conditions": [
                ("Temperature Range", "20 - 35 C"),
                ("Humidity Range", "60 - 70%"),
                ("Annual Rainfall", "500 - 1,200 mm"),
                ("Suitable Soil Types", "Loamy, Black, Red"),
                ("Growing Season", "Kharif (June - November)"),
                ("Crop Duration", "150 days"),
                ("Water Requirement", "Moderate"),
            ],
            "economics": [
                ("Yield per Acre", "1.5 - 3.0 tonnes"),
                ("Market Price", "Rs 8,000 - 15,000 per quintal"),
                ("MSP", "Not applicable"),
                ("Input Cost per Acre", "Rs 50,000"),
            ],
            "disease_link": "Not covered by the disease detection model.",
        },
    ]

    for crop in CROP_PROFILES:
        doc.add_page_break()
        h2 = doc.add_heading(f"{crop['number']} {crop['name']}", level=2)
        h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

        p = doc.add_paragraph()
        r = p.add_run(f"Category: {crop['icon']}")
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_paragraph(crop["description"])

        # Growing Conditions table
        p = doc.add_paragraph()
        r = p.add_run("Growing Conditions:")
        r.bold = True
        r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

        table = doc.add_table(rows=len(crop["conditions"]) + 1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_table_borders(table)
        table.rows[0].cells[0].text = "Parameter"
        table.rows[0].cells[1].text = "Requirement"
        header_row(table, 0)
        for i, (param, val) in enumerate(crop["conditions"]):
            table.rows[i + 1].cells[0].text = param
            table.rows[i + 1].cells[1].text = val

        doc.add_paragraph("")

        # Economics table
        p = doc.add_paragraph()
        r = p.add_run("Economics & Market Data:")
        r.bold = True
        r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

        table = doc.add_table(rows=len(crop["economics"]) + 1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_table_borders(table)
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Value"
        header_row(table, 0)
        for i, (metric, val) in enumerate(crop["economics"]):
            table.rows[i + 1].cells[0].text = metric
            table.rows[i + 1].cells[1].text = val

        doc.add_paragraph("")

        # Disease detection link
        p = doc.add_paragraph()
        r = p.add_run("Disease Detection Coverage: ")
        r.bold = True
        r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        p.add_run(crop["disease_link"])

    # ================================================================
    # 5. SUITABILITY SCORING ALGORITHM
    # ================================================================
    doc.add_page_break()
    h = doc.add_heading("5. Suitability Scoring Algorithm", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "Each crop is scored on a 0-100 scale using five weighted factors. "
        "The scoring algorithm evaluates how well the farmer's current conditions "
        "match each crop's ideal growing requirements."
    )

    # Scoring factors table
    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    scoring_data = [
        ("Factor", "Max Points", "Full Score Condition", "Partial Score"),
        ("Season Match", "25", "Current season matches crop's season", "5 pts if off-season"),
        ("Temperature", "25", "Current temp within crop's optimal range", "12 pts if within 5 C of range"),
        ("Humidity", "20", "Current humidity within crop's range", "10 pts if within 10% of range"),
        ("Rainfall", "15", "Annual rainfall within crop's range", "8 pts if >= 70% of minimum"),
        ("Soil Type", "15", "Farmer's soil matches crop's ideal soils", "5 pts for non-ideal soil"),
    ]
    for i, row_data in enumerate(scoring_data):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
        if i == 0:
            header_row(table, 0)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    r = p.add_run("Score Interpretation:")
    r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    score_ranges = [
        ("75 - 100: Highly Suitable", "Excellent match. Crop is strongly recommended for the given conditions. Displayed with green indicator."),
        ("50 - 74: Moderately Suitable", "Reasonable match with some limitations. May require additional inputs (irrigation, soil amendment). Displayed with orange indicator."),
        ("20 - 49: Marginally Suitable", "Possible but challenging. Significant adjustments needed. Displayed with red indicator."),
        ("Below 20: Not Recommended", "Conditions are too unfavorable. Crop is excluded from recommendations."),
    ]
    for title, detail in score_ranges:
        p = doc.add_paragraph()
        r = p.add_run(title + " - ")
        r.bold = True
        p.add_run(detail)

    # ================================================================
    # 6. SEASON CLASSIFICATION
    # ================================================================
    h = doc.add_heading("6. Season Classification", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "The system uses the Indian agricultural season classification. "
        "Season is auto-detected from the current month but can be manually overridden by the farmer."
    )

    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    season_data = [
        ("Season", "Months", "Characteristics"),
        ("Kharif", "June - October (Months 6-10)", "Monsoon/rainy season. Warm and humid. Major crops: Rice, Cotton, Maize, Soybean, Chilli."),
        ("Rabi", "November - March (Months 11-3)", "Winter season. Cool and dry. Major crops: Wheat, Potato, Mustard, Onion."),
        ("Zaid", "April - May (Months 4-5)", "Summer season. Hot and dry. Short-duration crops. All crops considered eligible."),
    ]
    for i, row_data in enumerate(season_data):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
        if i == 0:
            header_row(table, 0)

    # ================================================================
    # 7. SOIL TYPES SUPPORTED
    # ================================================================
    doc.add_page_break()
    h = doc.add_heading("7. Soil Types Supported", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "The system supports 7 major soil types found across India. The farmer selects "
        "their soil type, which is matched against each crop's ideal soil list."
    )

    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    soil_data = [
        ("Soil Type", "Characteristics", "Best Suited Crops"),
        ("Alluvial", "Most fertile, found in Indo-Gangetic plains. Rich in minerals.", "Rice, Wheat, Sugarcane, Potato, Maize, Mustard, Onion"),
        ("Black (Regur)", "Rich in clay, high moisture retention. Found in Deccan Plateau.", "Cotton, Soybean, Chilli"),
        ("Red", "Iron-rich, well-drained. Found in Southern and Eastern India.", "Tomato, Maize, Chilli"),
        ("Laterite", "Acidic, iron/aluminum oxides. Found in heavy rainfall areas.", "Limited crop suitability in this database"),
        ("Sandy", "Well-drained, low fertility. Found in Rajasthan, coastal areas.", "Potato, Bell Pepper, Mustard, Onion, Tomato"),
        ("Loamy", "Ideal mix of sand, silt, clay. Most versatile for agriculture.", "All 12 crops (most versatile soil)"),
        ("Clay", "High moisture retention, heavy texture. Expands when wet.", "Rice, Wheat, Sugarcane"),
    ]
    for i, row_data in enumerate(soil_data):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
        if i == 0:
            header_row(table, 0)

    # ================================================================
    # 8. REVENUE & PROFIT ESTIMATION
    # ================================================================
    h = doc.add_heading("8. Revenue & Profit Estimation", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "The system calculates estimated revenue and profit for each recommended crop "
        "based on the farmer's land area. The formulas used are:"
    )

    formulas = [
        ("Yield (tonnes)", "yield_per_acre x land_area_in_acres"),
        ("Revenue (Rs)", "yield_in_tonnes x 10 (quintals/tonne) x price_per_quintal"),
        ("Total Input Cost (Rs)", "cost_per_acre x land_area_in_acres"),
        ("Net Profit (Rs)", "Revenue - Total Input Cost"),
    ]

    p = doc.add_paragraph()
    r = p.add_run("Calculation Formulas:")
    r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    table = doc.add_table(rows=len(formulas) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Formula"
    header_row(table, 0)
    for i, (metric, formula) in enumerate(formulas):
        table.rows[i + 1].cells[0].text = metric
        r = table.rows[i + 1].cells[1].paragraphs[0].add_run(formula)
        r.font.name = "Consolas"
        r.font.size = Pt(9)

    doc.add_paragraph("")

    doc.add_paragraph(
        "Note: Revenue estimates use both minimum and maximum price ranges to provide "
        "the farmer with a best-case and worst-case scenario. The system also displays "
        "MSP (Minimum Support Price) for crops that have government price guarantees "
        "(Rice, Wheat, Cotton, Sugarcane, Maize, Soybean, Mustard)."
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    r = p.add_run("Land Area Conversion:")
    r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph("The system accepts land area in three units:")
    conversions = [
        "Acres (base unit, used in all calculations)",
        "Hectares (1 Hectare = 2.471 Acres)",
        "Bigha (1 Bigha = 0.625 Acres, using standard conversion)",
    ]
    for c in conversions:
        doc.add_paragraph(c, style="List Bullet")

    # ── Footer ──
    doc.add_paragraph("")
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = line.add_run("_" * 50)
    r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(
        "Generated for AI Agriculture Project - AGRI-X AI\n"
        "Crop Advisor Module | 12 Crops | 7 Soil Types | 3 Seasons\n"
        "Disease Detection: Tomato (10), Potato (3), Bell Pepper (2) = 15 classes"
    )
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"Word document saved successfully to: {OUTPUT_PATH}")


if __name__ == "__main__":
    create_document()
