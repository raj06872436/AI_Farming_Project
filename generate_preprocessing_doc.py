"""
Generate a Word document (.docx) documenting all preprocessing steps
used in the AI Farming Project for plant disease detection.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Preprocessing_Steps.docx"
)


def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "4472C4")
        borders.append(border)
    tblPr.append(borders)


def create_document():
    doc = Document()

    # ── Page Margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ── Styles ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Agriculture Project")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Plant Disease Detection Using Transfer Learning")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph("")

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("─" * 50)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph("")

    doc_title = doc.add_paragraph()
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_title.add_run("DATA PREPROCESSING STEPS")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph("")

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        "Comprehensive documentation of all data preprocessing,\n"
        "augmentation, and feature engineering techniques applied\n"
        "to the PlantVillage dataset for multi-class plant disease classification."
    )
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS
    # ================================================================
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    toc_items = [
        "1. Introduction & Overview",
        "2. Dataset Description",
        "3. Image Loading & Resizing",
        "4. Pixel Normalization (Rescaling)",
        "5. Data Splitting Strategy",
        "6. Data Augmentation Techniques",
        "7. Class Imbalance Handling",
        "8. Batch Processing & Generator Pipeline",
        "9. Inference-Time Preprocessing",
        "10. Preprocessing Pipeline Summary",
        "11. Configuration Parameters Reference",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_page_break()

    # ================================================================
    # 1. INTRODUCTION
    # ================================================================
    h = doc.add_heading("1. Introduction & Overview", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "This document describes the complete data preprocessing pipeline used in the "
        "AI Agriculture Project for plant disease detection. The project employs deep "
        "learning models based on transfer learning (MobileNetV2, ResNet50, EfficientNetB0, "
        "DenseNet121) to classify images of plant leaves into 15 disease/healthy categories."
    )
    doc.add_paragraph(
        "Preprocessing is a critical stage that transforms raw leaf images from the "
        "PlantVillage dataset into a standardized format suitable for training convolutional "
        "neural networks (CNNs). Proper preprocessing ensures consistent input dimensions, "
        "normalized pixel values, and augmented training diversity — all of which directly "
        "impact model accuracy and generalization."
    )

    p = doc.add_paragraph()
    run = p.add_run("Key Objectives of Preprocessing:")
    run.bold = True
    run.font.size = Pt(11)

    objectives = [
        "Standardize image dimensions to a uniform size accepted by pre-trained CNNs.",
        "Normalize pixel values to the [0, 1] range for stable gradient computation.",
        "Augment training data to improve model robustness and prevent overfitting.",
        "Handle class imbalance through computed class weights.",
        "Create efficient data pipelines using batched generators for memory-efficient training.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Bullet")

    # ================================================================
    # 2. DATASET DESCRIPTION
    # ================================================================
    h = doc.add_heading("2. Dataset Description", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "The project uses the PlantVillage dataset, a widely-used benchmark dataset for "
        "plant disease classification. It contains labeled images of plant leaves from "
        "three crop types: Pepper (Bell), Potato, and Tomato."
    )

    # Dataset summary table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    header_data = [
        ("Property", "Details"),
        ("Dataset Source", "PlantVillage (Open-source benchmark)"),
        ("Number of Classes", "15 (diseases + healthy)"),
        ("Image Format", "JPEG / PNG (RGB, 3 channels)"),
        ("Original Resolution", "256×256 pixels (variable)"),
    ]

    for i, (col1, col2) in enumerate(header_data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "1B5E20")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True

    doc.add_paragraph("")

    # Class names table
    p = doc.add_paragraph()
    run = p.add_run("Target Classes (15):")
    run.bold = True
    run.font.size = Pt(11)

    class_names = [
        ("1", "Pepper Bell — Bacterial Spot", "Disease"),
        ("2", "Pepper Bell — Healthy", "Healthy"),
        ("3", "Potato — Early Blight", "Disease"),
        ("4", "Potato — Late Blight", "Disease"),
        ("5", "Potato — Healthy", "Healthy"),
        ("6", "Tomato — Bacterial Spot", "Disease"),
        ("7", "Tomato — Early Blight", "Disease"),
        ("8", "Tomato — Late Blight", "Disease"),
        ("9", "Tomato — Leaf Mold", "Disease"),
        ("10", "Tomato — Septoria Leaf Spot", "Disease"),
        ("11", "Tomato — Spider Mites (Two-Spotted)", "Disease"),
        ("12", "Tomato — Target Spot", "Disease"),
        ("13", "Tomato — Yellow Leaf Curl Virus", "Disease"),
        ("14", "Tomato — Mosaic Virus", "Disease"),
        ("15", "Tomato — Healthy", "Healthy"),
    ]

    table = doc.add_table(rows=len(class_names) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    headers = ["#", "Class Name", "Type"]
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        set_cell_shading(cell, "1B5E20")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

    for i, (num, name, typ) in enumerate(class_names):
        row = table.rows[i + 1]
        row.cells[0].text = num
        row.cells[1].text = name
        row.cells[2].text = typ
        if typ == "Healthy":
            set_cell_shading(row.cells[2], "E8F5E9")

    # ================================================================
    # 3. IMAGE LOADING & RESIZING
    # ================================================================
    doc.add_page_break()
    h = doc.add_heading("3. Image Loading & Resizing", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "All images are loaded using Keras' ImageDataGenerator.flow_from_directory() method, "
        "which reads images directly from the folder structure where each subdirectory name "
        "represents a class label."
    )

    h2 = doc.add_heading("3.1 Resizing", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph(
        "All images are resized to 224 × 224 pixels using bilinear interpolation. "
        "This is the standard input size for pre-trained ImageNet models such as "
        "MobileNetV2, ResNet50, EfficientNetB0, and DenseNet121."
    )

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    resize_data = [
        ("Parameter", "Value"),
        ("Target Size", "224 × 224 pixels"),
        ("Color Space", "RGB (3 channels)"),
        ("Interpolation", "Bilinear (default in Keras)"),
    ]
    for i, (col1, col2) in enumerate(resize_data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                set_cell_shading(cell, "1B5E20")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True

    doc.add_paragraph("")

    h2 = doc.add_heading("3.2 Color Channel Format", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph(
        "Images are loaded in RGB format with a shape of (224, 224, 3). No color space "
        "conversion (e.g., to grayscale or HSV) is performed, as the pre-trained models "
        "expect 3-channel RGB input matching their ImageNet pre-training."
    )

    # Code snippet
    p = doc.add_paragraph()
    run = p.add_run("Code Reference:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    code = doc.add_paragraph()
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)
    run = code.add_run(
        'img_size = (224, 224)\n'
        'train_datagen.flow_from_directory(\n'
        '    data_dir,\n'
        '    target_size=img_size,\n'
        '    class_mode="categorical"\n'
        ')'
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x90)

    # ================================================================
    # 4. PIXEL NORMALIZATION
    # ================================================================
    h = doc.add_heading("4. Pixel Normalization (Rescaling)", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "All pixel values are rescaled from the original integer range [0, 255] to the "
        "floating-point range [0.0, 1.0] by dividing by 255. This normalization is applied "
        "uniformly to all data splits — training, validation, and test."
    )

    p = doc.add_paragraph()
    run = p.add_run("Why Normalize?")
    run.bold = True
    run.font.size = Pt(11)

    reasons = [
        "Ensures numerical stability during gradient-based optimization.",
        "Prevents large pixel values from dominating the loss function.",
        "Matches the expected input range of ImageNet pre-trained weights.",
        "Accelerates convergence during training.",
    ]
    for r in reasons:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_paragraph("")

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    norm_data = [
        ("Parameter", "Value"),
        ("Rescale Factor", "1.0 / 255"),
        ("Input Range", "[0, 255] (uint8)"),
        ("Output Range", "[0.0, 1.0] (float32)"),
    ]
    for i, (col1, col2) in enumerate(norm_data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                set_cell_shading(cell, "1B5E20")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Code Reference:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    code = doc.add_paragraph()
    run = code.add_run(
        '# Applied to ALL generators (train, validation, test)\n'
        'train_datagen = ImageDataGenerator(rescale=1.0 / 255, ...)\n'
        'val_test_datagen = ImageDataGenerator(rescale=1.0 / 255, ...)'
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x90)

    # ================================================================
    # 5. DATA SPLITTING STRATEGY
    # ================================================================
    doc.add_page_break()
    h = doc.add_heading("5. Data Splitting Strategy", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "The dataset is split into training and validation subsets using Keras' "
        "ImageDataGenerator with a validation_split parameter. The split is performed "
        "consistently across all training runs using a fixed random seed (seed=42) "
        "to ensure reproducibility."
    )

    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    split_data = [
        ("Split", "Proportion", "Purpose"),
        ("Training", "70%", "Model weight optimization with augmentation"),
        ("Validation", "15%", "Hyperparameter tuning & early stopping"),
        ("Test", "15%", "Final unbiased performance evaluation"),
        ("Seed", "42", "Ensures reproducible, deterministic splits"),
    ]
    for i, row_data in enumerate(split_data):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
        if i == 0:
            for cell in table.rows[i].cells:
                set_cell_shading(cell, "1B5E20")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True

    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: In the generator-based pipeline, the validation_split parameter "
        "of ImageDataGenerator is used to create the 80/20 split (training vs. validation+test). "
        "The configured ratios are 70/15/15, but in practice the generator-based approach "
        "uses an 80/20 split with the validation subset also serving as the test set. "
        "For rigorous evaluation (e.g., K-Fold cross-validation), the full dataset is "
        "loaded into memory and split with sklearn."
    )

    h2 = doc.add_heading("5.1 Cross-Validation", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph(
        "For advanced statistical analysis, Stratified K-Fold cross-validation "
        "(K=5) is employed. The full dataset is loaded into numpy arrays, and each "
        "fold maintains the same class distribution as the overall dataset."
    )

    # ================================================================
    # 6. DATA AUGMENTATION
    # ================================================================
    h = doc.add_heading("6. Data Augmentation Techniques", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "Data augmentation is applied ONLY to the training set to artificially increase "
        "its diversity and improve model generalization. The validation and test sets "
        "receive NO augmentation — only rescaling — to ensure unbiased evaluation."
    )

    p = doc.add_paragraph()
    run = p.add_run("Purpose of Data Augmentation:")
    run.bold = True

    aug_purposes = [
        "Prevents overfitting by exposing the model to varied versions of the same image.",
        "Simulates real-world variability (different lighting, camera angles, leaf orientations).",
        "Effectively increases the training set size without collecting new data.",
        "Improves model robustness to variations encountered in field conditions.",
    ]
    for ap in aug_purposes:
        doc.add_paragraph(ap, style="List Bullet")

    doc.add_paragraph("")

    # Augmentation parameters table
    table = doc.add_table(rows=10, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    aug_data = [
        ("Augmentation Technique", "Parameter Value", "Description"),
        ("Rotation Range", "30°", "Randomly rotates image up to ±30 degrees"),
        ("Horizontal Flip", "True", "Randomly mirrors image left-to-right"),
        ("Vertical Flip", "True", "Randomly mirrors image top-to-bottom"),
        ("Zoom Range", "0.2 (±20%)", "Random zoom in/out by up to 20%"),
        ("Brightness Range", "(0.8, 1.2)", "Random brightness adjustment ±20%"),
        ("Width Shift Range", "0.15 (15%)", "Horizontal translation up to 15% of width"),
        ("Height Shift Range", "0.15 (15%)", "Vertical translation up to 15% of height"),
        ("Shear Range", "0.15", "Random shearing transformation at 15%"),
        ("Fill Mode", "nearest", "Fills empty pixels with nearest pixel value"),
    ]
    for i, row_data in enumerate(aug_data):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
        if i == 0:
            for cell in table.rows[i].cells:
                set_cell_shading(cell, "1B5E20")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Important: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    run = p.add_run(
        "Augmentation is applied on-the-fly during training (real-time augmentation). "
        "Each epoch, the training images are randomly transformed differently, so the "
        "model never sees the exact same image twice. This is more memory-efficient than "
        "pre-generating augmented copies."
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Code Reference:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    code = doc.add_paragraph()
    run = code.add_run(
        'train_datagen = ImageDataGenerator(\n'
        '    rescale=1.0 / 255,\n'
        '    rotation_range=30,\n'
        '    horizontal_flip=True,\n'
        '    vertical_flip=True,\n'
        '    zoom_range=0.2,\n'
        '    brightness_range=(0.8, 1.2),\n'
        '    width_shift_range=0.15,\n'
        '    height_shift_range=0.15,\n'
        '    shear_range=0.15,\n'
        '    fill_mode="nearest",\n'
        '    validation_split=0.2\n'
        ')'
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x90)

    # ================================================================
    # 7. CLASS IMBALANCE HANDLING
    # ================================================================
    doc.add_page_break()
    h = doc.add_heading("7. Class Imbalance Handling", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "The PlantVillage dataset has unequal numbers of images across the 15 classes. "
        "If left unaddressed, the model may become biased toward majority classes. "
        "To counter this, class weights are computed inversely proportional to class "
        "frequency and applied during training."
    )

    h2 = doc.add_heading("7.1 Class Weight Computation", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph(
        "The weight for each class is calculated using the formula:"
    )

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run("weight(c) = total_samples / (num_classes × count(c))")
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Where total_samples is the total number of images, num_classes is 15, "
        "and count(c) is the number of images in class c. Classes with fewer images "
        "receive higher weights, making the loss function penalize misclassifications "
        "of underrepresented classes more heavily."
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Code Reference:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    code = doc.add_paragraph()
    run = code.add_run(
        'class_weights = {}\n'
        'for idx, (cls_name, count) in enumerate(sorted(class_counts.items())):\n'
        '    weight = total_samples / (n_classes * count)\n'
        '    class_weights[idx] = weight\n'
        '\n'
        '# Applied during model training:\n'
        'model.fit(train_data, ..., class_weight=class_weights)'
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x90)

    # ================================================================
    # 8. BATCH PROCESSING & GENERATOR PIPELINE
    # ================================================================
    h = doc.add_heading("8. Batch Processing & Generator Pipeline", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "The preprocessing pipeline uses Keras' ImageDataGenerator to create "
        "memory-efficient data generators that load and preprocess images in batches, "
        "rather than loading the entire dataset into RAM."
    )

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    batch_data = [
        ("Parameter", "Value"),
        ("Batch Size", "32 images"),
        ("Shuffle (Training)", "True — randomizes order each epoch"),
        ("Shuffle (Validation/Test)", "False — deterministic evaluation"),
        ("Class Mode", "categorical (one-hot encoded labels)"),
    ]
    for i, (col1, col2) in enumerate(batch_data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                set_cell_shading(cell, "1B5E20")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True

    doc.add_paragraph("")

    h2 = doc.add_heading("8.1 Label Encoding", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph(
        'Labels are encoded using one-hot encoding (class_mode="categorical"). '
        "Each label is represented as a vector of length 15 where only the index "
        "corresponding to the true class is set to 1. For example, if the true class "
        'is "Potato___Early_blight" (index 2), the label vector would be '
        "[0, 0, 1, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]."
    )

    h2 = doc.add_heading("8.2 Generator Pipeline Flow", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    steps = [
        ("Step 1: Directory Scan", "Scans the PlantVillage/ directory and maps subfolder names to class indices."),
        ("Step 2: Image Loading", "Reads each image file (JPEG/PNG) from disk into memory as a PIL Image."),
        ("Step 3: Resizing", "Resizes the image to 224×224 pixels using bilinear interpolation."),
        ("Step 4: Array Conversion", "Converts the PIL Image to a NumPy array of shape (224, 224, 3) with dtype float32."),
        ("Step 5: Rescaling", "Divides all pixel values by 255 to normalize to [0.0, 1.0]."),
        ("Step 6: Augmentation", "Applies random geometric and photometric transformations (training only)."),
        ("Step 7: Batching", "Groups images into batches of 32 for efficient GPU/CPU processing."),
        ("Step 8: Label Encoding", "Assigns one-hot encoded categorical labels to each image in the batch."),
    ]
    for step_title, step_desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(step_title + ": ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        run = p.add_run(step_desc)

    # ================================================================
    # 9. INFERENCE-TIME PREPROCESSING
    # ================================================================
    doc.add_page_break()
    h = doc.add_heading("9. Inference-Time Preprocessing", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "During inference (prediction on new images), the same preprocessing steps "
        "are applied as during training — EXCEPT for data augmentation. The inference "
        "preprocessing pipeline ensures that the model receives input in the exact "
        "same format it was trained on."
    )

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    inf_data = [
        ("Step", "Details"),
        ("1. Load Image", "Read image using Keras image.load_img()"),
        ("2. Resize", "Resize to 224×224 pixels (target_size)"),
        ("3. Convert to Array", "Convert to NumPy array using image.img_to_array()"),
        ("4. Rescale", "Divide by 255.0 to normalize to [0, 1]"),
    ]
    for i, (col1, col2) in enumerate(inf_data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                set_cell_shading(cell, "1B5E20")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True

    doc.add_paragraph("")

    doc.add_paragraph(
        "Additionally, a batch dimension is added using np.expand_dims(img_array, axis=0) "
        "to convert the shape from (224, 224, 3) to (1, 224, 224, 3), as the model "
        "expects batched input."
    )

    p = doc.add_paragraph()
    run = p.add_run("Code Reference (predict.py):")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    code = doc.add_paragraph()
    run = code.add_run(
        'img = image.load_img("leaf.jpg", target_size=(224, 224))\n'
        'img_array = image.img_to_array(img) / 255.0\n'
        'img_array = np.expand_dims(img_array, axis=0)\n'
        'prediction = model.predict(img_array)'
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x90)

    # ================================================================
    # 10. PIPELINE SUMMARY
    # ================================================================
    h = doc.add_heading("10. Preprocessing Pipeline Summary", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "The following table summarizes the complete preprocessing pipeline, "
        "showing which steps apply to each data split:"
    )

    table = doc.add_table(rows=9, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    summary_data = [
        ("Preprocessing Step", "Training", "Validation", "Test"),
        ("Image Loading (from directory)", "✓", "✓", "✓"),
        ("Resizing to 224×224", "✓", "✓", "✓"),
        ("Pixel Rescaling (÷255)", "✓", "✓", "✓"),
        ("Data Augmentation", "✓", "✗", "✗"),
        ("Class Weight Application", "✓", "✗", "✗"),
        ("Shuffle", "✓", "✗", "✗"),
        ("Batch Size = 32", "✓", "✓", "✓"),
        ("One-Hot Label Encoding", "✓", "✓", "✓"),
    ]
    for i, row_data in enumerate(summary_data):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if val == "✓":
                for paragraph in table.rows[i].cells[j].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
                        run.font.bold = True
            elif val == "✗":
                for paragraph in table.rows[i].cells[j].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        if i == 0:
            for cell in table.rows[i].cells:
                set_cell_shading(cell, "1B5E20")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True

    # ================================================================
    # 11. CONFIGURATION PARAMETERS REFERENCE
    # ================================================================
    doc.add_page_break()
    h = doc.add_heading("11. Configuration Parameters Reference", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    doc.add_paragraph(
        "All preprocessing parameters are centralized in the project's configuration "
        "system (src/config/settings.py). The following table provides a complete "
        "reference of all configurable preprocessing parameters:"
    )

    table = doc.add_table(rows=15, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    config_data = [
        ("Parameter", "Default Value", "Source"),
        ("dataset_path", "PlantVillage", "DATASET_PATH env var"),
        ("image_size", "224", "IMAGE_SIZE env var"),
        ("batch_size", "32", "BATCH_SIZE env var"),
        ("num_classes", "15", "NUM_CLASSES env var"),
        ("train_split", "0.70", "Hardcoded"),
        ("val_split", "0.15", "Hardcoded"),
        ("test_split", "0.15", "Hardcoded"),
        ("seed", "42", "SEED env var"),
        ("rotation_range", "30.0°", "Hardcoded"),
        ("horizontal_flip", "True", "Hardcoded"),
        ("zoom_range", "0.2", "Hardcoded"),
        ("brightness_range", "(0.8, 1.2)", "Hardcoded"),
        ("width_shift_range / height_shift_range", "0.15", "Hardcoded"),
    ]
    for i, row_data in enumerate(config_data):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
        if i == 0:
            for cell in table.rows[i].cells:
                set_cell_shading(cell, "1B5E20")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True

    doc.add_paragraph("")

    # ── Footer Note ──
    doc.add_paragraph("")
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("─" * 50)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "Generated for AI Agriculture Project — Plant Disease Detection\n"
        "Framework: TensorFlow / Keras | Dataset: PlantVillage\n"
        "Models: MobileNetV2, ResNet50, EfficientNetB0, DenseNet121"
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"\nWord document saved successfully to: {OUTPUT_PATH}")


if __name__ == "__main__":
    create_document()
