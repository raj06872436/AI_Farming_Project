"""
AGRI-X AI — Soil Dataset Parameters Documentation Generator
Generates a comprehensive Word document (.docx) listing all soil parameters
used across the project: SoilGrids API data, texture classification,
crop suitability scoring, yield prediction, and soil type definitions.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime


# ── Helper functions ──

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False,
                            color=None, size=None, alignment=None, space_after=None):
    """Add a paragraph with custom formatting."""
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    if alignment is not None:
        para.alignment = alignment
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    return para


def create_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table with alternating row colors."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2E7D32')

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(str(value))
            run.font.size = Pt(8.5)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'E8F5E9')

    # Apply column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def build_document():
    """Build the comprehensive soil parameters Word document."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    # ══════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AGRI-X AI')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(46, 125, 50)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Soil Dataset Parameters\nComprehensive Documentation')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f'Agricultural Intelligence Platform\n'
        f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n'
        f'Version: 2.0 — Phase 2'
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.italic = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction & Overview',
        '2. Soil Data Source — ISRIC SoilGrids API',
        '3. Core Soil Parameters (SoilGrids)',
        '   3.1 Clay Content',
        '   3.2 Sand Content',
        '   3.3 Silt Content',
        '   3.4 Soil pH (pH in H₂O)',
        '   3.5 Soil Organic Carbon (SOC)',
        '4. Soil Texture Classification System',
        '5. Soil Types Used in Crop Recommendation',
        '6. Crop–Soil Compatibility Matrix',
        '7. Soil Quality in Yield Prediction',
        '8. Soil Scoring in Crop Suitability Engine',
        '9. Parameters Summary Table',
        '10. Data Flow Architecture',
        '11. References',
    ]
    for item in toc_items:
        para = doc.add_paragraph(item)
        para.paragraph_format.space_after = Pt(2)
        run = para.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(33, 100, 50)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('1. Introduction & Overview', level=1)
    doc.add_paragraph(
        'The AGRI-X AI Agricultural Intelligence Platform uses soil data as a critical input '
        'for crop recommendation, yield prediction, and agricultural advisory services. '
        'Soil parameters are sourced from the ISRIC SoilGrids v2.0 REST API based on the '
        'user\'s geographic coordinates (latitude/longitude), and are used across multiple '
        'modules of the platform.'
    )
    doc.add_paragraph(
        'This document provides a comprehensive reference of every soil parameter used in '
        'the project, including its definition, data source, unit of measurement, value range, '
        'and how it is used in the platform\'s algorithms.'
    )

    doc.add_heading('Modules Using Soil Data', level=2)
    modules = [
        ('Location Service (location_service.py)',
         'Fetches raw soil data from ISRIC SoilGrids, classifies soil texture, '
         'and attaches soil properties to the user\'s location profile.'),
        ('Crop Advisor (crop_advisor.py)',
         'Uses soil type for crop suitability scoring. Each crop has a list of '
         'compatible soil types, and soil match contributes 15 out of 100 points '
         'to the crop suitability score.'),
        ('Yield Predictor (yield_predictor.py)',
         'Uses a soil quality factor (0.1–1.0) that contributes 15% of the '
         'environmental suitability factor for yield estimation.'),
        ('Insights Engine (insights_engine.py)',
         'Generates soil-quality-aware farming suggestions and fertilizer '
         'recommendations based on season and conditions.'),
        ('Farmer Chatbot (farmer_chatbot.py)',
         'Provides soil-related advice including soil testing, amendments, '
         'and crop rotation for soil health.'),
    ]
    for module_name, description in modules:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(module_name + ': ')
        run.bold = True
        run.font.size = Pt(10)
        para.add_run(description).font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. DATA SOURCE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('2. Soil Data Source — ISRIC SoilGrids API', level=1)
    doc.add_paragraph(
        'The platform retrieves soil composition data from ISRIC SoilGrids v2.0, a global '
        'gridded soil information system that provides predictions of soil properties at '
        '250-meter resolution worldwide.'
    )

    doc.add_heading('API Configuration', level=2)
    api_rows = [
        ['Endpoint', 'https://rest.isric.org/soilgrids/v2.0/properties/query'],
        ['Method', 'HTTP GET'],
        ['Queried Properties', 'clay, sand, silt, phh2o, soc'],
        ['Depth Layer', '0–5 cm (topsoil)'],
        ['Statistic', 'Mean (average prediction)'],
        ['Response Format', 'JSON'],
        ['Timeout', '10 seconds'],
        ['Caching', 'Streamlit @st.cache_data, TTL = 86400 s (24 hours)'],
        ['Resolution', '250 m × 250 m global grid'],
        ['Coverage', 'Global (land areas)'],
    ]
    table = doc.add_table(rows=len(api_rows) + 1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Parameter', 'Value']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1B5E20')

    for idx, (param, val) in enumerate(api_rows):
        row = table.rows[idx + 1]
        row.cells[0].text = ''
        run = row.cells[0].paragraphs[0].add_run(param)
        run.bold = True
        run.font.size = Pt(9)
        row.cells[1].text = ''
        run = row.cells[1].paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        if idx % 2 == 1:
            set_cell_shading(row.cells[0], 'E8F5E9')
            set_cell_shading(row.cells[1], 'E8F5E9')

    doc.add_paragraph()
    doc.add_paragraph(
        'Request Parameters:\n'
        '  • lon: Longitude of the location\n'
        '  • lat: Latitude of the location\n'
        '  • property: "clay,sand,silt,phh2o,soc"\n'
        '  • depth: "0-5cm"\n'
        '  • value: "mean"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. CORE SOIL PARAMETERS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('3. Core Soil Parameters (SoilGrids)', level=1)
    doc.add_paragraph(
        'The following five parameters are retrieved from the ISRIC SoilGrids API and form '
        'the foundation of all soil-based computations in the platform.'
    )

    # 3.1 Clay
    doc.add_heading('3.1 Clay Content', level=2)
    doc.add_paragraph(
        'Clay content represents the proportion of soil particles smaller than 0.002 mm '
        'in diameter. High clay content indicates heavy, water-retentive soil with slow '
        'drainage and good nutrient-holding capacity.'
    )
    clay_data = [
        ['API Property Name', 'clay'],
        ['Raw Unit (SoilGrids)', 'g/kg (grams per kilogram)'],
        ['Converted Unit (Platform)', '% (percentage, divided by 10)'],
        ['Depth Layer', '0–5 cm'],
        ['Typical Range', '5–60%'],
        ['Role in Classification', 'Clay > 40% → "Clay" soil type'],
        ['Agricultural Significance', 'High water retention, nutrient-rich, poor drainage, '
         'suits rice and sugarcane'],
    ]
    create_styled_table(doc, ['Attribute', 'Details'],
                        [[r[0], r[1]] for r in clay_data])

    doc.add_paragraph()

    # 3.2 Sand
    doc.add_heading('3.2 Sand Content', level=2)
    doc.add_paragraph(
        'Sand content represents soil particles between 0.05 mm and 2 mm in diameter. '
        'Sandy soils have excellent drainage but poor nutrient and water retention.'
    )
    sand_data = [
        ['API Property Name', 'sand'],
        ['Raw Unit (SoilGrids)', 'g/kg (grams per kilogram)'],
        ['Converted Unit (Platform)', '% (percentage, divided by 10)'],
        ['Depth Layer', '0–5 cm'],
        ['Typical Range', '10–90%'],
        ['Role in Classification', 'Sand > 70% → "Sandy" soil type\n'
         'Sand > 50% → "Sandy Loam"'],
        ['Agricultural Significance', 'Good drainage, quick warming, low nutrient retention, '
         'suits root crops like potato and onion'],
    ]
    create_styled_table(doc, ['Attribute', 'Details'],
                        [[r[0], r[1]] for r in sand_data])

    doc.add_paragraph()

    # 3.3 Silt
    doc.add_heading('3.3 Silt Content', level=2)
    doc.add_paragraph(
        'Silt content represents soil particles between 0.002 mm and 0.05 mm. '
        'Silt-rich soils are fertile and have moderate drainage with good moisture retention.'
    )
    silt_data = [
        ['API Property Name', 'silt'],
        ['Raw Unit (SoilGrids)', 'g/kg (grams per kilogram)'],
        ['Converted Unit (Platform)', '% (percentage, divided by 10)'],
        ['Depth Layer', '0–5 cm'],
        ['Typical Range', '5–70%'],
        ['Role in Classification', 'Silt > 60% → "Silty" soil type\n'
         'Silt > 40% → "Silt Loam"'],
        ['Agricultural Significance', 'Fertile, good moisture retention, moderate drainage, '
         'suits most crops'],
    ]
    create_styled_table(doc, ['Attribute', 'Details'],
                        [[r[0], r[1]] for r in silt_data])

    doc.add_paragraph()

    # 3.4 pH
    doc.add_heading('3.4 Soil pH (pH in H₂O)', level=2)
    doc.add_paragraph(
        'Soil pH measures the acidity or alkalinity of the soil on a scale of 0–14. '
        'It directly affects nutrient availability to plants. Most crops prefer a pH of '
        '6.0–7.5 (slightly acidic to neutral).'
    )
    ph_data = [
        ['API Property Name', 'phh2o'],
        ['Raw Unit (SoilGrids)', 'pH × 10 (integer)'],
        ['Converted Unit (Platform)', 'pH (divided by 10)'],
        ['Depth Layer', '0–5 cm'],
        ['Typical Range', '4.0–9.0'],
        ['Optimal Range (Platform)', '6.0–7.5 (shown as "Optimal" in green)'],
        ['Acceptable Range', '5.5–8.0 (shown as "Acceptable" in yellow)'],
        ['Extreme Range', '< 5.5 or > 8.0 (shown as "Extreme" in red)'],
        ['Display in UI', 'Numeric value with color-coded label (Optimal / Acceptable / Extreme)'],
    ]
    create_styled_table(doc, ['Attribute', 'Details'],
                        [[r[0], r[1]] for r in ph_data])

    doc.add_paragraph()

    # 3.5 SOC
    doc.add_heading('3.5 Soil Organic Carbon (SOC)', level=2)
    doc.add_paragraph(
        'Soil Organic Carbon measures the amount of carbon stored in soil organic matter. '
        'It is a key indicator of soil health and fertility, affecting water holding capacity, '
        'nutrient cycling, and soil structure.'
    )
    soc_data = [
        ['API Property Name', 'soc'],
        ['Raw Unit (SoilGrids)', 'dg/kg (decigrams per kilogram)'],
        ['Converted Unit (Platform)', 'g/kg (divided by 10)'],
        ['Depth Layer', '0–5 cm'],
        ['Typical Range', '1–100 g/kg'],
        ['Display Label', '"ORG. C" in the Terrain & Soil tab'],
        ['Agricultural Significance', 'Higher SOC = better soil fertility, water retention, '
         'and microbial activity'],
    ]
    create_styled_table(doc, ['Attribute', 'Details'],
                        [[r[0], r[1]] for r in soc_data])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. SOIL TEXTURE CLASSIFICATION
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('4. Soil Texture Classification System', level=1)
    doc.add_paragraph(
        'The platform classifies soil into texture types based on the proportions of clay, '
        'sand, and silt retrieved from SoilGrids. This classification is implemented in the '
        '_classify_soil() function in location_service.py.'
    )

    doc.add_heading('Classification Rules (Priority Order)', level=2)
    classification_rows = [
        ['1', 'Clay > 40%', 'Clay', 'Heavy soil, high water retention, poor drainage'],
        ['2', 'Sand > 70%', 'Sandy', 'Light soil, excellent drainage, poor nutrient retention'],
        ['3', 'Silt > 60%', 'Silty', 'Fertile, moderate drainage, good moisture retention'],
        ['4', 'Clay > 25% AND Sand > 25%', 'Clay Loam', 'Balanced heavy soil, moderate drainage'],
        ['5', 'Sand > 50%', 'Sandy Loam', 'Lightweight, good drainage, moderate fertility'],
        ['6', 'Silt > 40%', 'Silt Loam', 'Fertile, good structure, moderate drainage'],
        ['7', 'Default (none of above)', 'Loam', 'Ideal balanced soil — best for most crops'],
    ]
    create_styled_table(
        doc,
        ['Priority', 'Condition', 'Classification', 'Description'],
        classification_rows,
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: The classification follows a top-down priority order — the first matching '
        'condition determines the soil type. If no specific condition is met, the soil '
        'defaults to "Loam", which is considered ideal for most agricultural purposes.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. SOIL TYPES IN CROP RECOMMENDATION
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('5. Soil Types Used in Crop Recommendation', level=1)
    doc.add_paragraph(
        'The crop recommendation engine defines 7 standard soil types that are used for '
        'crop–soil compatibility matching. These are defined in the SOIL_TYPES list in '
        'crop_advisor.py.'
    )

    soil_types_detail = [
        ['Alluvial', 'Deposited by rivers. Very fertile, rich in minerals. '
         'Found in river plains and deltas.',
         'Clay, Silt, Sand mix', 'Rice, Wheat, Sugarcane, Maize, Potato, Mustard, Onion'],
        ['Black', 'Also called "Regur" or black cotton soil. High clay content, '
         'excellent moisture retention.',
         'High Clay (40-60%)', 'Cotton, Soybean, Chilli'],
        ['Red', 'Iron-rich, well-drained. Slightly acidic. Found in tropical regions.',
         'Sandy-Clay mix', 'Tomato, Maize, Chilli'],
        ['Laterite', 'Leached iron/aluminum-rich soil. Low fertility without amendments.',
         'Iron/Aluminum Oxides', 'Limited — requires heavy amendment'],
        ['Sandy', 'Coarse-grained, excellent drainage. Low nutrient holding capacity.',
         'Sand > 70%', 'Tomato, Potato, Bell Pepper, Mustard, Onion'],
        ['Loamy', 'Balanced mix of clay, silt, and sand. Ideal for most crops.',
         'Balanced mix', 'Rice, Wheat, Tomato, Potato, Bell Pepper, Cotton, Sugarcane, '
         'Maize, Soybean, Mustard, Onion, Chilli'],
        ['Clay', 'Heavy soil with > 40% clay. High moisture retention but poor drainage.',
         'Clay > 40%', 'Rice, Wheat, Sugarcane'],
    ]
    create_styled_table(
        doc,
        ['Soil Type', 'Description', 'Composition', 'Suitable Crops'],
        soil_types_detail,
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. CROP–SOIL COMPATIBILITY MATRIX
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('6. Crop–Soil Compatibility Matrix', level=1)
    doc.add_paragraph(
        'Each crop in the platform\'s database specifies a list of compatible soil types. '
        'During crop suitability scoring, a perfect soil match earns 15 points (out of 100 total), '
        'while a non-matching soil earns only 5 points. The following table shows the '
        'complete crop–soil compatibility data from CROP_DB in crop_advisor.py.'
    )

    crop_soil_matrix = [
        ['Rice', 'Alluvial, Clay, Loamy', '15 pts', '5 pts'],
        ['Wheat', 'Alluvial, Loamy, Clay', '15 pts', '5 pts'],
        ['Tomato', 'Loamy, Sandy, Red', '15 pts', '5 pts'],
        ['Potato', 'Loamy, Sandy, Alluvial', '15 pts', '5 pts'],
        ['Bell Pepper', 'Loamy, Sandy', '15 pts', '5 pts'],
        ['Cotton', 'Black, Alluvial', '15 pts', '5 pts'],
        ['Sugarcane', 'Alluvial, Loamy, Clay', '15 pts', '5 pts'],
        ['Maize', 'Loamy, Alluvial, Red', '15 pts', '5 pts'],
        ['Soybean', 'Black, Loamy', '15 pts', '5 pts'],
        ['Mustard', 'Loamy, Sandy, Alluvial', '15 pts', '5 pts'],
        ['Onion', 'Loamy, Sandy, Alluvial', '15 pts', '5 pts'],
        ['Chilli', 'Loamy, Black, Red', '15 pts', '5 pts'],
    ]
    create_styled_table(
        doc,
        ['Crop', 'Compatible Soil Types', 'Match Score', 'Non-Match Score'],
        crop_soil_matrix,
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. SOIL QUALITY IN YIELD PREDICTION
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('7. Soil Quality in Yield Prediction', level=1)
    doc.add_paragraph(
        'The Yield Prediction Engine (yield_predictor.py) uses a numerical "soil quality" '
        'factor that contributes to the overall environmental suitability score. This factor '
        'is a user-adjustable slider in the UI.'
    )

    doc.add_heading('Soil Quality Factor', level=2)
    yield_soil_rows = [
        ['Parameter Name', 'soil_quality'],
        ['Data Type', 'Float (continuous)'],
        ['Range', '0.1 – 1.0'],
        ['Default Value', '0.8 (Good soil)'],
        ['UI Control', 'Slider (step = 0.05)'],
        ['Weight in Env. Factor', '15% (0.15 weight)'],
        ['Formula', 'env_factor = temp_factor × 0.35 + humid_factor × 0.25 + '
         'rain_factor × 0.25 + soil_quality × 0.15'],
    ]
    create_styled_table(doc, ['Attribute', 'Details'],
                        [[r[0], r[1]] for r in yield_soil_rows])

    doc.add_paragraph()

    doc.add_heading('Soil Quality Interpretation', level=2)
    quality_rows = [
        ['0.1 – 0.3', 'Poor', 'Degraded soil, low organic matter, heavily compacted or saline'],
        ['0.3 – 0.5', 'Below Average', 'Low fertility, needs significant amendment'],
        ['0.5 – 0.7', 'Average', 'Moderate fertility, standard agricultural soil'],
        ['0.7 – 0.9', 'Good', 'Fertile, well-structured soil with adequate organic matter'],
        ['0.9 – 1.0', 'Excellent', 'Highly fertile, rich organic matter, ideal for all crops'],
    ]
    create_styled_table(
        doc,
        ['Range', 'Quality Level', 'Description'],
        quality_rows,
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'When soil quality is below 0.6, the yield predictor generates a recommendation: '
        '"Soil quality is low. Apply organic matter, compost, and soil amendments."'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 8. SOIL SCORING IN CROP SUITABILITY
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('8. Soil Scoring in Crop Suitability Engine', level=1)
    doc.add_paragraph(
        'The crop recommendation engine (_score_crop() in crop_advisor.py) uses a '
        'multi-factor scoring system out of 100 total points. Soil type matching is one '
        'of the five scoring components.'
    )

    doc.add_heading('Scoring Breakdown', level=2)
    scoring_rows = [
        ['Season Match', '25', 'Current season matches crop\'s optimal season'],
        ['Temperature', '25', 'Current temp within crop\'s optimal temperature range'],
        ['Humidity', '20', 'Current humidity within crop\'s optimal humidity range'],
        ['Rainfall', '15', 'Estimated annual rainfall matches crop needs'],
        ['Soil Type', '15', 'User-selected soil type matches crop\'s compatible soils'],
    ]
    create_styled_table(
        doc,
        ['Factor', 'Max Points', 'Description'],
        scoring_rows,
    )

    doc.add_paragraph()

    doc.add_heading('Soil Scoring Logic', level=2)
    doc.add_paragraph(
        'The soil scoring algorithm works as follows:'
    )
    soil_scoring_detail = [
        'If the selected soil type is in the crop\'s compatible soil list → +15 points '
        '(with message "✅ {soil} soil is ideal")',
        'If the selected soil type is NOT in the compatible list → +5 points '
        '(with message "⚠️ {soil} soil — not ideal but possible")',
        'The soil type is auto-filled from the location\'s detected soil type (via SoilGrids), '
        'but users can override it via the dropdown selector.',
    ]
    for detail in soil_scoring_detail:
        doc.add_paragraph(detail, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 9. PARAMETERS SUMMARY TABLE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('9. Complete Parameters Summary Table', level=1)
    doc.add_paragraph(
        'The following table consolidates all soil-related parameters used across the '
        'entire AGRI-X AI platform.'
    )

    summary_rows = [
        ['Clay Content', '%', '0–100', 'SoilGrids API', 'location_service.py',
         'Soil texture classification; displayed in UI'],
        ['Sand Content', '%', '0–100', 'SoilGrids API', 'location_service.py',
         'Soil texture classification; displayed in UI'],
        ['Silt Content', '%', '0–100', 'SoilGrids API', 'location_service.py',
         'Soil texture classification; displayed in UI'],
        ['pH (in H₂O)', 'pH', '4.0–9.0', 'SoilGrids API', 'location_service.py',
         'Soil acidity assessment; displayed with color-coded label'],
        ['Organic Carbon', 'g/kg', '1–100', 'SoilGrids API', 'location_service.py',
         'Soil health indicator; displayed in UI'],
        ['Soil Type (classified)', 'Category', '7 types', 'Derived', 'location_service.py',
         'Auto-classified from clay/sand/silt ratios'],
        ['Soil Type (crop matching)', 'Category', '7 types', 'User / Auto', 'crop_advisor.py',
         'Selected via dropdown; auto-filled from detected type'],
        ['Soil Quality', 'Float', '0.1–1.0', 'User Input', 'yield_predictor.py',
         'Adjustable slider; 15% weight in yield env. factor'],
        ['Soil Match Score', 'Points', '5 or 15', 'Computed', 'crop_advisor.py',
         'Part of 100-point crop suitability score'],
        ['Elevation', 'meters', '0–8848', 'Open-Meteo API', 'location_service.py',
         'Terrain context for soil interpretation'],
    ]
    create_styled_table(
        doc,
        ['Parameter', 'Unit', 'Range', 'Source', 'Module', 'Usage'],
        summary_rows,
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 10. DATA FLOW ARCHITECTURE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('10. Data Flow Architecture', level=1)
    doc.add_paragraph(
        'The soil data flows through the platform in the following sequence:'
    )

    flow_steps = [
        ('Step 1: Location Detection',
         'User\'s location is detected via IP geolocation (ip-api.com) or entered '
         'manually. Latitude and longitude coordinates are extracted.'),
        ('Step 2: SoilGrids API Query',
         'The get_soil_data() function queries ISRIC SoilGrids v2.0 with the coordinates. '
         'It requests clay, sand, silt, phh2o, and soc properties for the 0–5 cm depth layer.'),
        ('Step 3: Unit Conversion',
         'Raw values are converted: clay/sand/silt from g/kg to %, pH from pH×10 to pH, '
         'SOC from dg/kg to g/kg.'),
        ('Step 4: Soil Classification',
         '_classify_soil() determines the soil texture type (Clay, Sandy, Loam, etc.) '
         'based on the priority rules using clay, sand, and silt percentages.'),
        ('Step 5: Location Profile Enrichment',
         'Soil data, soil type, and elevation are attached to the location dictionary '
         'in session state for use by other modules.'),
        ('Step 6: Crop Advisor Usage',
         'The detected soil type auto-fills the soil selector dropdown. During scoring, '
         'the soil type is compared against each crop\'s compatible soils list.'),
        ('Step 7: Yield Predictor Usage',
         'The soil quality slider (default 0.8) feeds into the environmental factor '
         'calculation with a 15% weight.'),
        ('Step 8: UI Display',
         'The Terrain & Soil tab displays all raw soil properties (clay%, sand%, silt%, '
         'pH, organic carbon) in a styled card with the classified soil type.'),
    ]
    for step_title, step_desc in flow_steps:
        para = doc.add_paragraph()
        run = para.add_run(step_title + '\n')
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(46, 125, 50)
        run2 = para.add_run(step_desc)
        run2.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 11. REFERENCES
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('11. References', level=1)
    references = [
        'ISRIC SoilGrids v2.0 — https://rest.isric.org/soilgrids/v2.0/ — Global gridded '
        'soil information at 250m resolution.',
        'ISRIC World Soil Information — https://www.isric.org/ — International Soil Reference '
        'and Information Centre.',
        'Open-Meteo Elevation API — https://open-meteo.com/ — Free elevation data API.',
        'USDA Soil Texture Triangle — Standard classification of soil texture based on '
        'clay, sand, and silt proportions.',
        'FAO Soil Classification — Food and Agriculture Organization soil taxonomy guidelines.',
        'Indian Council of Agricultural Research (ICAR) — Soil type classifications used '
        'in Indian agriculture (Alluvial, Black, Red, Laterite).',
    ]
    for ref in references:
        doc.add_paragraph(ref, style='List Number')

    doc.add_paragraph()

    # Footer note
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        '— End of Document —\n'
        f'Generated by AGRI-X AI Soil Parameters Documentation Tool\n'
        f'{datetime.now().strftime("%B %d, %Y")}'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

    return doc


if __name__ == "__main__":
    print("[*] Generating Soil Dataset Parameters documentation...")
    doc = build_document()
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "Soil_Dataset_Parameters.docx"
    )
    doc.save(output_path)
    print(f"[OK] Document saved: {output_path}")
    print(f"[i] File size: {os.path.getsize(output_path) / 1024:.1f} KB")
